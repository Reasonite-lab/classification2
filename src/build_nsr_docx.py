from __future__ import annotations

import re
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"
FIGURES = ROOT / "figures"

FIGURE_FILES = {
    1: FIGURES / "petdb_primary4_validation_performance.png",
    2: FIGURES / "petdb_primary4_permutation_importance.png",
    3: FIGURES / "cross_database_transfer.png",
    4: FIGURES / "multiclass_external_validation.png",
    5: FIGURES / "east_african_rift_transition.png",
}

FIGURE_TEXT = {
    "en": {
        1: (
            "Figure 1. Group-aware validation of the PetDB four-class task. "
            "Macro-F1 and balanced accuracy from five-fold out-of-fold predictions "
            "using random, citation-set, citation-overlap-component and first-order-"
            "location splits. The 21-variable panel contains only elements observed "
            "in all four PetDB classes.",
            "Alt text: Paired bar charts show high random-split performance for random "
            "forest and logistic regression, followed by lower citation-group "
            "performance and the lowest performance when locations are held out.",
        ),
        2: (
            "Figure 2. Conservative permutation importance. Mean decrease in balanced "
            "accuracy after permuting each variable within citation-overlap-held-out "
            "folds; error bars show between-fold standard deviation. Importance "
            "describes model dependence and is not a unique causal attribution.",
            "Alt text: A horizontal bar chart ranks TiO2 and Sr as the two strongest "
            "variables, followed by K2O, Zr, MnO, Nb and Y, with uncertainty bars "
            "across five held-out folds.",
        ),
        3: (
            "Figure 3. Bidirectional cross-database transfer. Three-class ARC-CIB-OIB "
            "transfer between GEOROC and PetDB using the same 21 variables. The target "
            "database was not used for fitting. Broad and rift-aligned CIB ontologies "
            "are compared.",
            "Alt text: Paired bar charts show stronger transfer from GEOROC to PetDB "
            "than from PetDB to GEOROC; aligning the continental intraplate class to "
            "rift volcanics improves the reverse direction but does not remove the gap.",
        ),
        4: (
            "Figure 4. Four-setting external transfer and back-arc pressure test. "
            "Row-normalized confusion matrices, class recall and unscored Vate Trough "
            "prediction shares for balanced logistic regression and balanced random "
            "forest. The scored test has 231 samples and two province proxies for each "
            "class.",
            "Alt text: Confusion matrices and recall bars show near-perfect arc and "
            "mid-ocean-ridge classification, strong ocean-island classification and "
            "weaker continental intraplate classification; Vate Trough predictions are "
            "dominated by mid-ocean-ridge affinity.",
        ),
        5: (
            "Figure 5. East African Rift as a rift-plume geochemical transition. The "
            "entire East African Rift was removed from training. Panel A summarizes "
            "mean four-class probabilities by country; panel B shows element medians "
            "shifting the database CIB cohort toward the OIB training domain.",
            "Alt text: Stacked country bars are dominated by ocean-island probability, "
            "while a horizontal ranking shows that SiO2, FeOT and Nb provide the largest "
            "median shifts from other continental intraplate basalts toward ocean-island "
            "basalts.",
        ),
    },
    "zh": {
        1: (
            "图1. PetDB四分类任务的分组验证。五折折外预测下，比较随机、引文集合、"
            "引文重叠连通组分和一级地点拆分的macro-F1与平衡准确率。21变量面板仅"
            "包含PetDB四类均有观测的元素。",
            "替代文本：成对柱状图显示随机拆分下随机森林和逻辑回归性能最高，引文"
            "分组后降低，地点留出时最低。",
        ),
        2: (
            "图2. 保守置换重要性。在每个引文重叠留出折中置换变量后，计算平衡准确率"
            "的平均下降；误差线为折间标准差。重要性描述模型依赖，不代表唯一因果归因。",
            "替代文本：水平柱状图把TiO2和Sr列为最强变量，随后为K2O、Zr、MnO、Nb"
            "和Y，并显示五个留出折之间的不确定性。",
        ),
        3: (
            "图3. 跨数据库双向迁移。使用相同21变量在GEOROC与PetDB间开展ARC-CIB-"
            "OIB三分类迁移。目标数据库从不参与拟合，并比较宽义和裂谷对齐的CIB本体。",
            "替代文本：成对柱状图显示GEOROC到PetDB强于PetDB到GEOROC；把大陆板内"
            "类别对齐为裂谷火山岩可以改善反向迁移，但不能消除差距。",
        ),
        4: (
            "图4. 四类外部迁移与弧后压力测试。类别平衡逻辑回归和随机森林的行归一化"
            "混淆矩阵、逐类召回率及不计分Vate Trough预测比例。计分测试含231件样品，"
            "每类2个省域代理。",
            "替代文本：混淆矩阵和召回率柱状图显示ARC与MORB接近完全正确，OIB较强"
            "而CIB较弱；Vate Trough预测以MORB亲和性为主。",
        ),
        5: (
            "图5. 东非裂谷作为裂谷-地幔柱地球化学过渡。训练时完整删除东非裂谷。"
            "A图给出各国四类平均概率；B图显示使数据库CIB队列向OIB训练域移动的元素"
            "中位数。",
            "替代文本：各国堆叠柱以OIB概率为主；水平排序显示SiO2、FeOT和Nb使东非"
            "裂谷CIB相对于其他CIB更接近OIB中位数。",
        ),
    },
}

TABLE_ROWS = {
    "en": [
        ["Class", "Province proxy", "Source", "n", "RF recall"],
        ["ARC", "South New Hebrides arc front", "PANGAEA.922011", "12", "1.000"],
        ["ARC", "Costa Rica volcanic front", "GEOROC 2JETOA", "18", "0.944"],
        ["CIB", "Rio Grande Rift/Jemez Lineament", "Rowe et al. Table S1", "28", "0.464"],
        ["CIB", "Big Pine volcanic field", "GEOROC 2JETOA", "70", "0.543"],
        ["MORB", "Southwest Indian Ridge", "Figshare 25295671", "36", "1.000"],
        ["MORB", "South Mid-Atlantic Ridge, 18.0-20.6 S", "Zhong et al. Table S1", "12", "1.000"],
        ["OIB", "Mauna Loa 2022 eruption", "Rhoads et al. Table S1", "16", "1.000"],
        ["OIB", "La Palma 2021 eruption", "Day et al. Table S1", "39", "0.718"],
    ],
    "zh": [
        ["类别", "省域代理", "来源", "n", "RF召回率"],
        ["ARC", "南新赫布里底弧前缘", "PANGAEA.922011", "12", "1.000"],
        ["ARC", "哥斯达黎加火山前缘", "GEOROC 2JETOA", "18", "0.944"],
        ["CIB", "Rio Grande Rift/Jemez Lineament", "Rowe等 Table S1", "28", "0.464"],
        ["CIB", "Big Pine火山场", "GEOROC 2JETOA", "70", "0.543"],
        ["MORB", "西南印度洋中脊", "Figshare 25295671", "36", "1.000"],
        ["MORB", "南大西洋中脊18.0-20.6 S", "Zhong等 Table S1", "12", "1.000"],
        ["OIB", "Mauna Loa 2022年喷发", "Rhoads等 Table S1", "16", "1.000"],
        ["OIB", "La Palma 2021年喷发", "Day等 Table S1", "39", "0.718"],
    ],
}

TABLE_CAPTION = {
    "en": (
        "Table 1. Strict external whole-rock basalt cohorts. The two-province "
        "design is a minimum symmetry check, not enough clusters for stable "
        "province-level confidence intervals. The Southwest Indian Ridge repository "
        "cohort remains provisional because a linked journal article was not "
        "established; the second MORB cohort is publisher-primary and publication-linked."
    ),
    "zh": (
        "表1. 严格外部全岩玄武岩队列。每类2个省域代理只达到最小对称设计，仍不足以"
        "估计稳定的省域聚类置信区间。西南印度洋中脊仓储队列仍属暂定，因为尚未确定"
        "关联期刊论文；第二个MORB队列来自出版商原始补充表并有明确论文关联。"
    ),
}


def set_run_font(run, language: str, size: float | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(
        qn("w:eastAsia"), "SimSun" if language == "zh" else "Times New Roman"
    )
    if size is not None:
        run.font.size = Pt(size)


def set_style_font(style, language: str, size: float, bold=False, italic=False) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "SimSun" if language == "zh" else "Times New Roman")


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, "en", 9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_document(doc: Document, language: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_page_field(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    set_style_font(normal, language, 12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    for name, size, bold, italic, before, after in [
        ("Heading 1", 12, True, False, 12, 6),
        ("Heading 2", 12, True, False, 10, 4),
        ("Heading 3", 12, False, True, 8, 3),
    ]:
        style = doc.styles[name]
        set_style_font(style, language, size, bold, italic)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = doc.styles["Caption"]
    set_style_font(caption, language, 9.5)
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)

    # Use Word's native numbered-list definition so references remain structurally
    # numbered when authors insert, delete, or reorder entries during revision.
    reference = doc.styles["List Number"]
    set_style_font(reference, language, 10.5)
    reference.paragraph_format.left_indent = Inches(0.28)
    reference.paragraph_format.first_line_indent = Inches(-0.28)
    reference.paragraph_format.line_spacing = 1.0
    reference.paragraph_format.space_after = Pt(4)


def add_inline(paragraph, text: str, language: str, size: float | None = None) -> None:
    tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for token in tokens:
        if not token:
            continue
        bold = token.startswith("**") and token.endswith("**")
        italic = (
            token.startswith("*")
            and token.endswith("*")
            and not token.startswith("**")
        )
        content = token[2:-2] if bold else token[1:-1] if italic else token
        run = paragraph.add_run(content)
        set_run_font(run, language, size)
        run.bold = bold
        run.italic = italic


def add_title_page(doc: Document, language: str) -> None:
    if language == "en":
        title = "Geochemical memory limits machine-learning discrimination of basalt tectonic settings"
        article_type = "RESEARCH ARTICLE"
        running = "Running title: Geochemical memory in basalt discrimination"
        author = "Bangkun Xu (author name and order to be confirmed)1,*"
        affiliation = "1 [Department, institution, city, postal code and country to be confirmed]"
        corresponding = (
            "*Corresponding author. E-mail: xubangkun439@gmail.com; telephone and fax: "
            "[to be confirmed]."
        )
        teaser = (
            "Teaser text: This study shows that basalt classifiers transfer across "
            "end-member settings but fail systematically where inherited subduction or "
            "plume signals decouple rock chemistry from present tectonic position."
        )
    else:
        title = "地球化学记忆限制了机器学习对玄武岩构造环境的判别"
        article_type = "研究论文（中文审阅翻译版）"
        running = "短标题：玄武岩判别中的地球化学记忆"
        author = "徐邦坤（英文署名及作者排序待确认）1,*"
        affiliation = "1 [院系、单位、城市、邮编和国家待确认]"
        corresponding = (
            "*通讯作者。电子邮箱：xubangkun439@gmail.com；电话及传真：[待确认]。"
        )
        teaser = (
            "推广短句：本研究表明，玄武岩分类器可以跨越端元构造环境迁移，但当继承的"
            "俯冲或地幔柱信号使岩石化学与现今构造位置解耦时，会发生系统性失效。"
        )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(article_type)
    set_run_font(r, language, 12)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run_font(r, language, 16)
    r.bold = True

    for value, size, after in [
        (running, 11, 26),
        (author, 12, 10),
        (affiliation, 11, 10),
        (corresponding, 10.5, 28),
        (teaser, 10.5, 10),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        add_inline(p, value, language, size)
    if language == "zh":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        add_inline(
            p,
            "本文件供科学内容审阅；正式向NSR投稿时使用英文版。",
            language,
            10.5,
        )
    doc.add_page_break()


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    table_width = table_pr.first_child_found_in("w:tblW")
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    indent = table_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)

    borders = table_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)
    for edge in ["top", "bottom", "insideH"]:
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:color"), "666666")
        borders.append(node)
    for edge in ["left", "right", "insideV"]:
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)


def add_table(doc: Document, language: str) -> None:
    caption = doc.add_paragraph(style="Caption")
    caption.paragraph_format.keep_with_next = True
    add_inline(caption, TABLE_CAPTION[language], language, 9.5)
    rows = TABLE_ROWS[language]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    header_properties.append(table_header)
    widths = [900, 3000, 2400, 800, 2260]
    for i, values in enumerate(rows):
        for j, value in enumerate(values):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if j in [0, 3, 4]
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, language, 8.5)
            r.bold = i == 0
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, number: int, language: str) -> None:
    path = FIGURE_FILES[number]
    with Image.open(path) as image:
        px_width, px_height = image.size
    ratio = px_width / px_height
    max_width = 6.35
    max_height = 6.35 if number == 2 else 6.6
    width = min(max_width, max_height * ratio)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    alt = FIGURE_TEXT[language][number][1].split(":", 1)[-1].strip()
    shape._inline.docPr.set("descr", alt)
    caption, alt_text = FIGURE_TEXT[language][number]
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.keep_with_next = True
    add_inline(p, caption, language, 9.5)
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    add_inline(p, alt_text, language, 9)


def is_reference(line: str) -> re.Match[str] | None:
    return re.match(r"^(\d+)\.\s+(.*)$", line)


def build(language: str) -> Path:
    source = PAPER / (
        "nsr_manuscript_v1.md" if language == "en" else "nsr_manuscript_v1_zh.md"
    )
    output = PAPER / f"NSR_manuscript_v1_{language}.docx"
    text = source.read_text(encoding="utf-8")
    start_marker = "## ABSTRACT" if language == "en" else "## 摘要"
    stop_marker = "## FIGURE LEGENDS" if language == "en" else "## 图注"
    body = text.split(start_marker, 1)[1].split(stop_marker, 1)[0]
    lines = [start_marker] + body.splitlines()

    doc = Document()
    configure_document(doc, language)
    doc.core_properties.title = (
        "Geochemical memory limits machine-learning discrimination of basalt tectonic settings"
        if language == "en"
        else "地球化学记忆限制了机器学习对玄武岩构造环境的判别"
    )
    doc.core_properties.author = "Bangkun Xu (provisional)"
    doc.core_properties.subject = "NSR Research Article draft v1"
    doc.core_properties.comments = "Generated from frozen project results; author metadata pending confirmation."
    add_title_page(doc, language)

    in_references = False
    skip_table_block = False
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        figure_match = re.fullmatch(
            r"\[INSERT FIGURE (\d+) HERE\]", line
        ) or re.fullmatch(r"\[此处插入图(\d+)\]", line)
        if figure_match:
            add_figure(doc, int(figure_match.group(1)), language)
            continue
        if line in {"[INSERT TABLE 1 HERE]", "[此处插入表1]"}:
            add_table(doc, language)
            skip_table_block = True
            continue
        if skip_table_block:
            if line.startswith("|") or line.startswith("**Table 1.") or line.startswith("**表1."):
                continue
            skip_table_block = False
        if line.startswith("## "):
            heading = line[3:]
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.page_break_before = heading in {
                "REFERENCES",
                "参考文献",
            }
            add_inline(p, heading, language, 12)
            in_references = heading in {"REFERENCES", "参考文献"}
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:], language, 12)
            continue
        if line.startswith("**Keywords:") or line.startswith("**关键词："):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            add_inline(p, line, language, 11)
            continue
        reference = is_reference(line) if in_references else None
        if reference:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, reference.group(2), language, 10.5)
            continue
        if line.startswith(">"):
            continue
        p = doc.add_paragraph()
        add_inline(p, line, language)

    doc.save(output)
    return output


def main() -> None:
    for language in ["en", "zh"]:
        output = build(language)
        print(output)


if __name__ == "__main__":
    main()
